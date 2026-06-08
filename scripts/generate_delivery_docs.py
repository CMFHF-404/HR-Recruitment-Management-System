from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "交付文档"
PACKAGE_DIR = OUT_DIR / "第4回_查询可视化与系统测试_交付包"
SHOT_DIR = PACKAGE_DIR / "screenshots"


def set_font(run, name="Microsoft YaHei", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    for name, size, color in [
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13, "2F5597"),
        ("Heading 3", 11.5, "1F4E79"),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def title(doc, text, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=18, bold=True, color="1F4E79")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_font(r, size=10.5, color="666666")
    doc.add_paragraph()


def shade_cell(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        shade_cell(hdr[i])
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_font(r, bold=True)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                for r in p.runs:
                    set_font(r, size=9.5)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return t


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def add_image(doc, filename, caption):
    path = SHOT_DIR / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(6.2))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_font(r, size=9, color="666666")


def callout(doc, heading, body, fill="EAF3F8"):
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(heading)
    set_font(r, size=10.5, bold=True, color="1F4E79")
    p.add_run("\n")
    r = p.add_run(body)
    set_font(r, size=10, color="333333")
    doc.add_paragraph()


def add_explained_image(doc, filename, caption, note):
    path = SHOT_DIR / filename
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(6.2))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        set_font(r, size=9, bold=True, color="666666")
    p = doc.add_paragraph()
    r = p.add_run("图片说明：")
    set_font(r, size=10, bold=True, color="1F4E79")
    r = p.add_run(note)
    set_font(r, size=10)
    doc.add_paragraph()


def build_lab4_doc():
    doc = Document()
    style_doc(doc)
    title(doc, "第4回 数据查询、可视化与系统测试结果交付文档", "小型企业招聘管理信息系统")

    doc.add_heading("一、实验定位与完成情况", level=1)
    doc.add_paragraph(
        "本次按讲义“自选系统路径”完成第4回交付内容，在既有 Spring Boot + Vue 招聘管理系统中补齐查询、统计可视化、CSV 导出、系统测试与交付说明。"
    )
    table(
        doc,
        ["讲义要求", "本项目完成情况"],
        [
            ["数据查询", "候选人支持姓名、电话、邮箱关键词查询，并支持按应聘岗位筛选；岗位支持岗位/部门关键词查询。"],
            ["结果展示", "候选人、岗位、招聘流程均以表格分页展示，候选人可查看完整招聘进度。"],
            ["统计可视化", "统计首页展示 5 个指标卡，并用 ECharts 柱状图展示各岗位招聘漏斗。"],
            ["地图或业务图表", "本系统非 GIS 项目，使用岗位维度业务统计图替代地图展示。"],
            ["CSV 导出", "新增候选人 CSV 导出接口和前端按钮，导出文件保存于 exports/candidates.csv。"],
            ["系统测试", "后端自动化测试覆盖登录、权限、异常输入、流程约束、统计和 CSV 导出；前端通过生产构建。"],
            ["交付材料", "已整理 README、requirements.txt、数据库脚本、导出 CSV、截图和交付包 zip。"],
        ],
        [1.4, 5.0],
    )

    doc.add_heading("二、运行环境与启动步骤", level=1)
    numbered(doc, "启动后端：进入 backend，执行 $env:SPRING_PROFILES_ACTIVE='h2'，再执行 .\\mvnw.cmd spring-boot:run。")
    numbered(doc, "启动前端：进入 frontend，执行 npm install，再执行 npm run dev -- --host 127.0.0.1 --port 5179。")
    numbered(doc, "浏览器访问 http://127.0.0.1:5179，HR 账号：admin / admin123，部门主管账号：manager / manager123。")
    numbered(doc, "执行后端测试：cd backend 后运行 .\\mvnw.cmd test。")
    numbered(doc, "执行前端构建：cd frontend 后运行 npm run build。")

    doc.add_heading("三、查询、可视化与导出功能", level=1)
    table(
        doc,
        ["功能", "入口", "说明"],
        [
            ["候选人查询", "候选人管理", "按姓名、电话、邮箱关键词检索；按岗位下拉筛选。"],
            ["岗位查询", "岗位管理", "按岗位名称或所属部门关键词检索。"],
            ["招聘流程查询", "招聘流程", "分页查看筛选、面试、录用记录，并维护流程状态。"],
            ["统计指标", "统计首页", "展示岗位数量、候选人数、待筛选、面试人数、录用人数。"],
            ["统计图表", "统计首页", "以柱状图对比各岗位候选人数、面试人数、录用人数。"],
            ["CSV 导出", "候选人管理", "导出候选人基础资料、应聘岗位、部门、简历附件和创建时间。"],
        ],
        [1.2, 1.4, 3.8],
    )

    doc.add_heading("四、关键截图", level=1)
    add_image(doc, "截图1_系统首页统计.png", "图1 统计首页与岗位招聘漏斗图")
    add_image(doc, "截图2_候选人查询与CSV导出.png", "图2 候选人查询与 CSV 导出入口")
    add_image(doc, "截图3_招聘流程测试页面.png", "图3 招聘流程维护与系统测试页面")
    add_image(doc, "截图4_岗位数据查询页面.png", "图4 岗位查询与岗位数据维护页面")

    doc.add_heading("五、测试与交付文件", level=1)
    table(
        doc,
        ["项目", "结果"],
        [
            ["后端 CSV 导出测试", "新增 hrCanExportCandidatesAsCsv 用例，验证 text/csv、表头、姓名、邮箱、岗位字段。"],
            ["后端完整测试", "覆盖登录失败、角色权限、主管确认、面试前置、录用前置、附件类型异常、AI 失败保护等。"],
            ["前端构建", "通过 npm run build 验证生产构建。"],
            ["README", "交付包 README.txt 说明系统主题、功能、数据库、运行命令和账号。"],
            ["requirements", "根目录与交付包均提供 requirements.txt，列出 Java、Node、MySQL、Spring Boot、Vue 等依赖。"],
            ["导出文件", "exports/candidates.csv。"],
        ],
        [1.8, 4.6],
    )

    doc.save(OUT_DIR / "第4回_查询可视化与系统测试_结果交付文档.docx")


def build_project_doc():
    doc = Document()
    style_doc(doc)
    title(doc, "小型企业招聘管理信息系统总项目说明", "过程文档任务与最终交付说明")

    table(
        doc,
        ["项目项", "说明"],
        [
            ["系统名称", "小型企业招聘管理信息系统"],
            ["建设目标", "将招聘流程中的岗位、候选人、筛选、主管确认、面试、录用和统计归档为一套可演示、可测试、可交付的信息系统。"],
            ["技术路线", "Spring Boot 3 + Vue 3 + Element Plus + ECharts + MySQL/H2"],
            ["默认账号", "HR：admin / admin123；部门主管：manager / manager123"],
            ["交付范围", "需求说明、系统分析与设计、数据字典、测试用例、README、requirements.txt、截图说明与第4回交付包。"],
        ],
        [1.3, 5.1],
    )
    callout(
        doc,
        "交付摘要",
        "本文档按课程过程文档任务整理，覆盖必做项中的需求说明、系统分析与设计、测试与交付三大部分，并补充系统运行截图及图片说明，便于答辩展示和归档检查。",
    )

    doc.add_heading("一、项目概述", level=1)
    doc.add_paragraph(
        "本系统面向小型企业招聘场景，覆盖岗位发布、候选人登记、简历筛选、部门主管二次确认、面试安排、录用结果登记和招聘数据统计。系统采用前后端分离架构，后端为 Spring Boot REST API，前端为 Vue 3 管理界面，数据库使用 MySQL，课堂演示可使用 H2。"
    )
    table(
        doc,
        ["模块", "业务价值", "主要输出"],
        [
            ["岗位管理", "沉淀部门招聘需求，避免岗位信息分散。", "岗位列表、岗位状态、岗位要求。"],
            ["候选人管理", "统一维护候选人资料与简历附件。", "候选人档案、简历文本、CSV 导出。"],
            ["招聘流程", "规范筛选、主管确认、面试和录用前置条件。", "筛选状态、面试记录、录用结果。"],
            ["统计分析", "支持答辩展示与招聘进展复盘。", "指标卡、岗位招聘漏斗图。"],
        ],
        [1.3, 2.8, 2.3],
    )

    doc.add_heading("二、需求说明", level=1)
    doc.add_heading("2.1 用户故事", level=2)
    stories = [
        "作为 HR，我希望登录系统后管理候选人信息，以便统一维护招聘资料。",
        "作为 HR，我希望按姓名、电话、邮箱搜索候选人，以便快速定位人员记录。",
        "作为 HR，我希望按岗位筛选候选人，以便查看某岗位招聘进展。",
        "作为 HR，我希望上传 PDF、DOCX 或 TXT 简历，以便沉淀候选人附件和简历文本。",
        "作为 HR，我希望查看 AI 匹配度和快评，以便辅助初筛判断。",
        "作为部门主管，我希望创建和维护岗位，以便提交部门招聘需求。",
        "作为部门主管，我希望确认 HR 初筛通过的候选人，以便控制进入面试的人选质量。",
        "作为 HR，我希望只有主管确认通过后才能安排面试，以便保证流程合规。",
        "作为 HR，我希望候选人完成面试后才能登记录用结果，以便避免绕过流程。",
        "作为管理者，我希望查看岗位维度统计图，以便了解候选、面试和录用转化情况。",
        "作为 HR，我希望导出候选人 CSV，以便用于课设归档、线下统计或提交材料。",
    ]
    for item in stories:
        bullet(doc, item)

    doc.add_heading("2.2 角色与权限", level=2)
    table(
        doc,
        ["角色", "账号示例", "主要权限"],
        [
            ["HR", "admin / admin123", "候选人管理、简历上传、筛选处理、面试安排、录用登记、统计查看、CSV 导出。"],
            ["部门主管", "manager / manager123", "岗位创建/编辑/关闭/删除，待确认候选人审核。"],
            ["未登录用户", "-", "仅可访问登录接口，其他 API 均拒绝。"],
        ],
        [1.1, 1.5, 3.8],
    )

    doc.add_heading("2.3 数据质量规则", level=2)
    rules = [
        "管理员用户名必须唯一，密码使用 BCrypt 加密保存。",
        "岗位名称、所属部门、招聘人数、岗位要求不能为空。",
        "岗位招聘人数必须为正整数，前端输入框最小值为 1。",
        "岗位状态仅允许 OPEN、CLOSED。",
        "候选人姓名、性别、电话、邮箱、学历、毕业院校和应聘岗位不能为空。",
        "候选人邮箱必须符合邮箱格式。",
        "候选人电话必须符合手机号或常见电话号码格式。",
        "候选人必须关联一个存在的岗位。",
        "候选人创建后自动生成筛选、面试和录用流程记录。",
        "简历附件仅支持 PDF、DOCX、TXT，且必须能提取有效文本。",
        "只有 HR 初筛通过并经主管确认通过后，才能安排或完成面试。",
        "只有面试完成后，才能登记非待定录用结果。",
        "岗位下已有候选人时禁止删除，可改为关闭状态。",
    ]
    for item in rules:
        bullet(doc, item)

    doc.add_heading("三、系统分析与设计", level=1)
    doc.add_heading("3.1 业务流程图（泳道简化版）", level=2)
    table(
        doc,
        ["阶段", "部门主管", "HR", "系统"],
        [
            ["岗位需求", "创建/维护岗位", "查看开放岗位", "保存岗位并校验必填字段"],
            ["候选登记", "-", "登记候选人、上传简历", "保存候选人并生成筛选/面试/录用记录"],
            ["简历初筛", "-", "查看 AI 快评并录入筛选状态", "保存筛选状态；通过后生成主管待办"],
            ["主管确认", "审核通过或驳回", "查看确认结果", "限制未确认候选人进入面试"],
            ["面试安排", "-", "安排时间、地点、面试官并记录评价", "校验主管确认结果"],
            ["录用登记", "-", "登记录用、未录用、放弃入职", "校验面试已完成"],
            ["统计交付", "查看岗位信息", "查看统计图、导出 CSV", "聚合统计并生成导出文件"],
        ],
        [1.0, 1.8, 1.9, 2.2],
    )

    doc.add_heading("3.2 功能结构图", level=2)
    for item in [
        "小型企业招聘管理信息系统",
        "  登录与认证：账号登录、JWT 鉴权、角色权限控制",
        "  岗位管理：岗位新增、编辑、关闭、删除、查询",
        "  候选人管理：候选人新增、编辑、删除、查询、岗位筛选、简历上传、CSV 导出",
        "  简历筛选：AI 分析、筛选状态、筛选意见",
        "  主管确认：待确认列表、通过、驳回、确认意见",
        "  面试管理：时间、地点、面试官、面试状态、评价",
        "  录用管理：待定、录用、未录用、放弃入职、薪资说明、备注",
        "  统计分析：总览指标、岗位招聘漏斗图、待筛选数量",
    ]:
        doc.add_paragraph(item)

    doc.add_heading("3.3 ER 图说明", level=2)
    doc.add_paragraph(
        "admin 为独立账号表；position 与 candidate 为一对多；candidate 与 resume_screening、interview、offer_result 均为一对一。删除候选人时同步删除其流程记录；删除岗位前需确认岗位下没有候选人。"
    )
    table(
        doc,
        ["实体关系", "关系说明"],
        [
            ["position 1 - N candidate", "一个岗位可对应多个候选人，一个候选人只能应聘一个岗位。"],
            ["candidate 1 - 1 resume_screening", "每个候选人对应一条简历筛选与主管确认记录。"],
            ["candidate 1 - 1 interview", "每个候选人对应一条面试安排与评价记录。"],
            ["candidate 1 - 1 offer_result", "每个候选人对应一条最终录用结果记录。"],
        ],
        [2.2, 4.2],
    )

    doc.add_heading("3.4 数据字典", level=2)
    data_rows = [
        ["admin.id", "账号编号", "BIGINT", "是", "自增主键"],
        ["admin.username", "用户名", "VARCHAR(50)", "是", "唯一"],
        ["admin.role", "角色", "VARCHAR(20)", "是", "HR / MANAGER"],
        ["position.name", "岗位名称", "VARCHAR(100)", "是", "不为空"],
        ["position.department", "所属部门", "VARCHAR(100)", "是", "不为空"],
        ["position.headcount", "招聘人数", "INT", "是", ">=1"],
        ["position.status", "岗位状态", "VARCHAR(20)", "是", "OPEN / CLOSED"],
        ["candidate.name", "候选人姓名", "VARCHAR(50)", "是", "不为空"],
        ["candidate.phone", "联系电话", "VARCHAR(30)", "是", "手机号或常见电话格式"],
        ["candidate.email", "邮箱", "VARCHAR(120)", "是", "邮箱格式"],
        ["candidate.position_id", "应聘岗位", "BIGINT", "是", "必须引用 position.id"],
        ["candidate.resume_original_file_name", "简历原文件名", "VARCHAR(255)", "否", "PDF/DOCX/TXT"],
        ["resume_screening.status", "筛选状态", "VARCHAR(20)", "是", "PENDING / PASSED / REJECTED"],
        ["resume_screening.ai_match_score", "AI 匹配度", "INT", "否", "0-100"],
        ["resume_screening.manager_status", "主管确认状态", "VARCHAR(20)", "是", "NOT_SUBMITTED / PENDING / APPROVED / REJECTED"],
        ["interview.status", "面试状态", "VARCHAR(30)", "是", "NOT_SCHEDULED / SCHEDULED / COMPLETED / CANCELED"],
        ["offer_result.status", "录用状态", "VARCHAR(30)", "是", "PENDING / OFFERED / REJECTED / ABANDONED"],
    ]
    table(doc, ["字段", "含义", "类型", "必填", "范围/规则"], data_rows, [1.8, 1.5, 1.2, 0.7, 2.1])

    doc.add_heading("四、测试与交付", level=1)
    doc.add_heading("4.1 测试用例", level=2)
    tests = [
        ["TC01", "HR 正确账号登录", "admin/admin123", "登录成功并返回 HR token"],
        ["TC02", "错误密码登录", "admin/wrong", "返回失败提示"],
        ["TC03", "部门主管登录", "manager/manager123", "登录成功并返回 MANAGER 角色"],
        ["TC04", "HR 越权创建岗位", "HR 调用 POST /api/positions", "返回 403"],
        ["TC05", "主管创建岗位", "合法岗位名称、部门、人数、要求", "保存成功"],
        ["TC06", "候选人邮箱异常", "email=abc", "表单或后端校验失败"],
        ["TC07", "候选人电话异常", "phone=abc", "后端返回格式不正确"],
        ["TC08", "上传不支持附件", "resume.png", "返回仅支持 PDF、DOCX、TXT"],
        ["TC09", "空文本简历", "空 TXT", "返回未能提取有效文本"],
        ["TC10", "未主管确认安排面试", "HR 初筛通过但主管未确认", "返回需经部门主管确认"],
        ["TC11", "未完成面试登记录用", "面试未完成直接 OFFERED", "返回需完成面试"],
        ["TC12", "CSV 导出", "HR 访问 /api/candidates/export", "返回 text/csv，包含表头、姓名、邮箱、岗位"],
        ["TC13", "统计首页", "完成录用流程后查询统计", "录用人数增加，各岗位图表有数据"],
    ]
    table(doc, ["编号", "测试项", "输入/操作", "预期结果"], tests, [0.7, 1.5, 2.1, 2.1])

    doc.add_heading("4.2 README 与运行步骤", level=2)
    for item in [
        "后端：cd backend；$env:SPRING_PROFILES_ACTIVE='h2'；.\\mvnw.cmd spring-boot:run。",
        "前端：cd frontend；npm install；npm run dev -- --host 127.0.0.1 --port 5179。",
        "访问：http://127.0.0.1:5179。",
        "后端测试：cd backend；.\\mvnw.cmd test。",
        "前端构建：cd frontend；npm run build。",
    ]:
        bullet(doc, item)

    doc.add_heading("4.3 requirements.txt 依赖列表", level=2)
    for item in [
        "运行环境：Java >= 17、Node.js >= 20、npm >= 10、MySQL >= 8.0。",
        "后端依赖：Spring Boot 3.5.7、Spring Web、Spring Data JPA、Spring Security、Validation、MySQL Connector/J、H2、PDFBox、POI、JUnit。",
        "前端依赖：Vue 3、Vite 8、Element Plus、Axios、ECharts、Vue Router、@element-plus/icons-vue。",
    ]:
        bullet(doc, item)

    doc.add_page_break()
    doc.add_heading("五、系统运行截图与图片说明", level=1)
    doc.add_paragraph(
        "本节图片来自第4回交付包 screenshots 目录，用于说明系统主要页面与交付功能的实际运行效果。"
    )
    add_explained_image(
        doc,
        "截图1_系统首页统计.png",
        "图1 统计首页与岗位招聘漏斗",
        "该图展示 HR 登录后的统计首页。上方指标卡汇总岗位数量、候选人数、待筛选人数、面试人数和录用人数；下方柱状图按岗位展示候选、面试、录用三类数据，满足第4回“统计可视化”和“非 GIS 项目业务图表”要求。",
    )
    add_explained_image(
        doc,
        "截图2_候选人查询与CSV导出.png",
        "图2 候选人查询与 CSV 导出",
        "该图展示候选人管理页面。页面顶部可按姓名、电话、邮箱关键词查询，并可按应聘岗位筛选；右侧“导出 CSV”按钮对应本次新增的候选人数据导出功能，导出的文件保存于交付包 exports/candidates.csv。",
    )
    add_explained_image(
        doc,
        "截图3_招聘流程测试页面.png",
        "图3 招聘流程维护页面",
        "该图展示招聘流程页面中的简历筛选、面试安排和录用结果三个页签。系统通过状态字段和后端校验控制流程顺序，例如未通过主管确认不能安排面试、未完成面试不能登记录用结果。",
    )
    add_explained_image(
        doc,
        "截图4_岗位数据查询页面.png",
        "图4 岗位查询与岗位数据维护",
        "该图展示岗位管理页面。部门主管可维护岗位名称、所属部门、招聘人数、岗位要求和岗位状态；HR 可查看开放岗位并在候选人登记时关联岗位，形成 position 与 candidate 的一对多关系。",
    )

    output_path = OUT_DIR / "小型企业招聘管理信息系统_总项目说明.docx"
    try:
        doc.save(output_path)
    except PermissionError:
        doc.save(OUT_DIR / "小型企业招聘管理信息系统_总项目说明_新版截图.docx")


if __name__ == "__main__":
    build_lab4_doc()
    build_project_doc()
